"""Readers for the document formats accepted by the application."""

import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document as DocxDocument
from loguru import logger

from document_analyzer.core.pdf_reader import PdfDocument, PdfPage, PdfReadError, read_pdf
from document_analyzer.core.settings import Settings, get_settings

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".doc", ".txt", ".md"})


def read_document(path: str | Path, settings: Settings | None = None) -> PdfDocument:
    """Extract a supported document into the common page-aware model."""

    document_path = Path(path)
    suffix = document_path.suffix.lower()
    logger.debug("Selected a reader for {} ({})", document_path, suffix or "[no extension]")
    if suffix not in SUPPORTED_EXTENSIONS:
        raise PdfReadError(f"Unsupported document type: {suffix or '[none]'}")
    if suffix == ".pdf":
        document = read_pdf(document_path, settings)
        logger.debug("PDF reader extracted {} pages from {}", len(document.pages), document_path)
        return document
    if suffix == ".docx":
        document = _read_docx(document_path, "docx")
        logger.debug("DOCX reader extracted {} sections from {}", len(document.pages), document_path)
        return document
    if suffix == ".doc":
        document = _read_doc(document_path, settings or get_settings())
        logger.debug("DOC reader extracted {} sections from {}", len(document.pages), document_path)
        return document
    document = _read_text(document_path, suffix[1:])
    logger.debug("Text reader extracted {} sections from {}", len(document.pages), document_path)
    return document


def document_from_extracted_pages(
    path: str | Path,
    filename: str,
    file_type: str,
    pages: list[str],
    metadata: dict[str, str] | None = None,
) -> PdfDocument:
    """Reconstruct a document from extraction persisted during upload.

    This deliberately does not touch the original file, so page refreshes and
    chat requests never repeat PDF parsing or OCR.
    """
    if not pages or not any(page.strip() for page in pages):
        raise PdfReadError(f"Document has no persisted extractable text: {filename}")
    return PdfDocument(
        path=Path(path),
        filename=filename,
        metadata=dict(metadata or {}),
        pages=tuple(PdfPage(number=index, text=text) for index, text in enumerate(pages, start=1)),
        file_type=file_type,
    )


def _read_docx(path: Path, file_type: str) -> PdfDocument:
    if not path.is_file():
        raise PdfReadError(f"Document does not exist: {path}")
    try:
        source = DocxDocument(str(path))
        sections: list[str] = [paragraph.text.strip() for paragraph in source.paragraphs if paragraph.text.strip()]
        for table in source.tables:
            sections.extend(" | ".join(cell.text.strip() for cell in row.cells).strip() for row in table.rows)
        text = "\n".join(section for section in sections if section)
    except Exception as exc:
        raise PdfReadError(f"Unable to read DOCX {path}: {exc}") from exc
    return _document_from_text(path, text, file_type)


def _read_doc(path: Path, settings: Settings) -> PdfDocument:
    logger.info("Started DOC conversion for {} using {}", path.name, settings.libreoffice_command)
    executable = shutil.which(settings.libreoffice_command)
    if executable is None:
        raise PdfReadError(f"Cannot read {path.name}: {settings.libreoffice_command!r} is required for .doc files")
    with tempfile.TemporaryDirectory(prefix="document-analyzer-doc-") as temp_dir:
        try:
            subprocess.run(
                [executable, "--headless", "--convert-to", "docx", "--outdir", temp_dir, str(path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=settings.conversion_timeout_seconds,
            )
            converted = Path(temp_dir) / f"{path.stem}.docx"
            if not converted.is_file():
                raise PdfReadError(f"LibreOffice did not create a DOCX for {path.name}")
            converted_document = _read_docx(converted, "doc")
        except PdfReadError:
            raise
        except Exception as exc:
            raise PdfReadError(f"Unable to convert DOC {path}: {exc}") from exc
    logger.info("Converted {} into {} sections", path.name, len(converted_document.pages))
    return PdfDocument(
        path=path,
        filename=path.name,
        metadata=converted_document.metadata,
        pages=converted_document.pages,
        file_type="doc",
    )


def _read_text(path: Path, file_type: str) -> PdfDocument:
    if not path.is_file():
        raise PdfReadError(f"Document does not exist: {path}")
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise PdfReadError(f"Document is not valid UTF-8: {path}") from exc
    except OSError as exc:
        raise PdfReadError(f"Unable to read document {path}: {exc}") from exc
    return _document_from_text(path, text, file_type)


def _document_from_text(path: Path, text: str, file_type: str) -> PdfDocument:
    if not text.strip():
        raise PdfReadError(f"Document contains no extractable text: {path}")
    stat = path.stat()
    metadata = {"filename": path.name, "path": str(path), "size_bytes": str(stat.st_size)}
    return PdfDocument(
        path=path,
        filename=path.name,
        metadata=metadata,
        pages=(PdfPage(number=1, text=text.strip()),),
        file_type=file_type,
    )
