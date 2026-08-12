from pathlib import Path

from docx import Document

from document_analyzer.core.document_reader import document_from_extracted_pages, read_document


def test_read_text_and_markdown_documents(tmp_path: Path) -> None:
    text_path = tmp_path / "notes.txt"
    text_path.write_text("plain text", encoding="utf-8")
    markdown_path = tmp_path / "notes.md"
    markdown_path.write_text("# heading\n\nbody", encoding="utf-8")

    assert read_document(text_path).pages[0].text == "plain text"
    assert "# heading" in read_document(markdown_path).content


def test_read_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "notes.docx"
    source = Document()
    source.add_paragraph("A paragraph")
    table = source.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "left"
    table.rows[0].cells[1].text = "right"
    source.save(path)

    document = read_document(path)

    assert "A paragraph" in document.content
    assert "left | right" in document.content


def test_document_from_extracted_pages_preserves_order_and_metadata(tmp_path: Path) -> None:
    document = document_from_extracted_pages(
        tmp_path / "retained.pdf",
        "original.pdf",
        "pdf",
        ["first page", "second page"],
        {"document_id": "document-123", "extraction_method": "ocr"},
    )

    assert [page.number for page in document.pages] == [1, 2]
    assert document.content == "first page\n\nsecond page"
    assert document.document_id == "document-123"
    assert document.metadata["extraction_method"] == "ocr"
